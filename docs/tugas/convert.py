import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_paragraph_with_runs(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    
    # Simple markdown inline parser (**bold**, *italic*)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)
    return p

def convert_md_to_docx(md_path, docx_path, title):
    doc = Document()
    
    # Configure document page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Document Title Page
    t_p = doc.add_paragraph()
    t_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = t_p.add_run(title.upper())
    t_run.bold = True
    t_run.font.size = Pt(22)
    t_run.font.color.rgb = RGBColor(12, 35, 64) # Navy Blue
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_text = []
    
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip('\n')
        stripped = line.strip()
        
        # 1. Code Block Handling
        if stripped.startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                # Shading / border XML for code box
                pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="CCCCCC"/></w:pBdr>')
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                p._p.get_or_add_pPr().append(pBdr)
                p._p.get_or_add_pPr().append(shd)
                
                run = p.add_run('\n'.join(code_text))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(50, 50, 50)
                code_text = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_text.append(line)
            i += 1
            continue
            
        # 2. Table Handling
        if stripped.startswith('|'):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            # We reached the end of the table
            in_table = False
            # Parse accumulated rows
            parsed_rows = []
            for r in table_rows:
                cells = [c.strip() for c in r.split('|')[1:-1]]
                # Skip divider row (e.g. |---|---|)
                if cells and all(re.match(r'^:?-+:?$', c) for c in cells):
                    continue
                parsed_rows.append(cells)
                
            if parsed_rows:
                num_cols = max(len(r) for r in parsed_rows)
                num_rows = len(parsed_rows)
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                
                for r_idx, row_data in enumerate(parsed_rows):
                    for c_idx, val in enumerate(row_data):
                        if c_idx < len(table.rows[r_idx].cells):
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            # Inline style parsing for cells
                            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', val)
                            for part in parts:
                                if part.startswith('**') and part.endswith('**'):
                                    run = p.add_run(part[2:-2])
                                    run.bold = True
                                elif part.startswith('*') and part.endswith('*'):
                                    run = p.add_run(part[1:-1])
                                    run.italic = True
                                else:
                                    p.add_run(part)
                            
                            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                            if r_idx == 0:
                                # Header background (Navy)
                                set_cell_background(cell, "0C2340")
                                for run in p.runs:
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                    run.bold = True
                            elif r_idx % 2 == 1:
                                # Zebra striping
                                set_cell_background(cell, "F2F5F8")
                                
                doc.add_paragraph() # space after table
            table_rows = []
            
        # 3. Headers
        if stripped.startswith('# '):
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[2:])
            run.font.color.rgb = RGBColor(12, 35, 64)
            run.bold = True
        elif stripped.startswith('## '):
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[3:])
            run.font.color.rgb = RGBColor(24, 76, 120)
            run.bold = True
        elif stripped.startswith('### '):
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped[4:])
            run.font.color.rgb = RGBColor(30, 100, 150)
            run.bold = True
        elif stripped.startswith('#### '):
            p = doc.add_heading(level=4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[5:])
            run.font.color.rgb = RGBColor(50, 50, 50)
            run.bold = True
        # 4. Lists
        elif stripped.startswith('- ') or stripped.startswith('* '):
            add_paragraph_with_runs(doc, stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', stripped):
            match = re.match(r'^\d+\.\s(.*)', stripped)
            add_paragraph_with_runs(doc, match.group(1), style='List Number')
        # 5. Empty lines
        elif not stripped:
            pass
        # 6. Standard Paragraphs
        else:
            add_paragraph_with_runs(doc, stripped)
            
        i += 1
        
    doc.save(docx_path)
    print(f"Successfully converted {md_path} -> {docx_path}")

if __name__ == '__main__':
    convert_md_to_docx('ComPro-SRS .md', 'ComPro-SRS.docx', 'Software Requirement Specification')
    convert_md_to_docx('ComPro-SDD.md', 'ComPro-SDD.docx', 'Software Design Document')
